import io
import logging
import os
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

import fitz
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Emu
from pdf2docx import Converter

WATERMARK_TOKENS = ("licensed", "drishti", "khanna", "mecstudio")
FAILED_PAGE_RE = re.compile(r"Ignore page (\d+) due to making page error")

_SOFFICE_CACHE = None


def find_soffice():
    global _SOFFICE_CACHE
    if _SOFFICE_CACHE is not None:
        return _SOFFICE_CACHE
    override = os.environ.get("SANITIZER_SOFFICE")
    if override and Path(override).exists():
        _SOFFICE_CACHE = override
        return _SOFFICE_CACHE
    for name in ("soffice", "soffice.exe"):
        found = shutil.which(name)
        if found:
            _SOFFICE_CACHE = found
            return _SOFFICE_CACHE
    candidates = []
    if os.name == "nt":
        program_dirs = [os.environ.get("ProgramFiles", r"C:\Program Files"),
                        os.environ.get("ProgramFiles(x86)", r"C:\Program Files (x86)"),
                        os.environ.get("LOCALAPPDATA", r"C:\Program Files")]
        for base in program_dirs:
            candidates.extend([Path(base) / "LibreOffice" / "program" / "soffice.exe"])
    elif sys.platform == "darwin":
        candidates.append(Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"))
    else:
        candidates.extend([Path("/usr/bin/soffice"), Path("/usr/local/bin/soffice"),
                           Path("/opt/libreoffice/program/soffice")])
    for cand in candidates:
        if cand.exists():
            _SOFFICE_CACHE = str(cand)
            return _SOFFICE_CACHE
    _SOFFICE_CACHE = ""
    return _SOFFICE_CACHE


class LibreOfficeNotFoundError(RuntimeError):
    pass


class _FailureCollector(logging.Handler):
    def __init__(self):
        super().__init__(level=logging.ERROR)
        self.pages = set()

    def emit(self, record):
        m = FAILED_PAGE_RE.search(record.getMessage())
        if m:
            self.pages.add(int(m.group(1)) - 1)


def _redact_watermark(pdf_path: Path) -> Path:
    doc = fitz.open(str(pdf_path))
    cleaned = [_redact_page(page) for page in doc]
    if any(cleaned):
        out = pdf_path.with_suffix(".clean.pdf")
        doc.save(str(out))
        doc.close()
        return out
    doc.close()
    return pdf_path


def _redact_page(page) -> bool:
    hits = [
        fitz.Rect(w[:4])
        for w in page.get_text("words")
        if any(t in w[4].lower() for t in WATERMARK_TOKENS)
    ]
    if not hits:
        return False
    box = hits[0]
    for r in hits[1:]:
        box |= r
    box.x0 -= 2
    box.y0 -= 2
    box.x1 += 2
    box.y1 += 2
    page.add_redact_annot(box)
    page.apply_redactions()
    return True


class MSWord:
    """MS Word COM automation context (Windows only)."""

    def __init__(self):
        import pythoncom
        import win32com.client

        self._pythoncom = pythoncom
        pythoncom.CoInitialize()
        self.app = win32com.client.DispatchEx("Word.Application")
        self.app.Visible = False
        self.app.DisplayAlerts = 0

    def __enter__(self):
        return self.app

    def __exit__(self, exc_type, exc, tb):
        try:
            self.app.Quit(0)
        except Exception:
            pass
        self._pythoncom.CoUninitialize()
        return False


def word_doc_to_docx(doc_path: Path, out_dir: Path) -> Path:
    out_path = (out_dir / (doc_path.stem + ".docx")).resolve()
    with MSWord() as word:
        doc = word.Documents.Open(str(doc_path.resolve()), ReadOnly=True)
        try:
            doc.SaveAs2(str(out_path), FileFormat=16)
        finally:
            doc.Close(0)
    if not out_path.exists():
        raise RuntimeError(f"MS Word failed to convert {doc_path.name}")
    return out_path


def word_docx_to_pdf(docx_path: Path, out_dir: Path) -> Path:
    out_path = (out_dir / (docx_path.stem + ".pdf")).resolve()
    with MSWord() as word:
        doc = word.Documents.Open(str(docx_path.resolve()), ReadOnly=True)
        try:
            doc.ExportAsFixedFormat(str(out_path), ExportFormat=17)
        finally:
            doc.Close(0)
    if not out_path.exists():
        raise RuntimeError(f"MS Word failed to export {docx_path.name} to PDF")
    return out_path


def pdf_to_docx(pdf_path: Path, out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    clean = _redact_watermark(pdf_path)
    temps = [clean] if clean != pdf_path else []
    try:
        portrait = _portraitize_pdf(clean)
        if portrait != clean:
            temps.append(portrait)
            clean = portrait

        out_path = out_dir / (pdf_path.stem + ".docx")
        collector = _FailureCollector()
        logger = logging.getLogger("pdf2docx")
        logger.addHandler(collector)
        try:
            cv = Converter(str(clean))
            try:
                cv.convert(str(out_path))
            finally:
                cv.close()
            if collector.pages:
                _patch_failed_pages(Path(clean), out_path, sorted(collector.pages))
        finally:
            logger.removeHandler(collector)
    finally:
        for t in temps:
            t.unlink(missing_ok=True)
    return out_path


def _page_content_width(doc, section_index: int) -> Emu:
    sections = list(doc.sections)
    sec = sections[min(section_index, len(sections) - 1)]
    return sec.page_width - sec.left_margin - sec.right_margin


def _patch_failed_pages(clean_pdf: Path, docx_path: Path, failed: list[int]):
    doc = Document(str(docx_path))
    body = doc.element.body
    sect_ps = body.findall(f".//{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}sectPr")
    total_sections = len(sect_ps)
    pdf = fitz.open(str(clean_pdf))
    successful = [p for p in range(len(pdf)) if p not in set(failed)]
    offset = 0
    for f in failed:
        n_before = sum(1 for s in successful if s < f)
        img = _page_image_para(doc, pdf, f, _page_content_width(doc, max(n_before - 1, 0)))
        brk = _page_break_para(doc)
        if n_before == 0 or total_sections == 0:
            anchor = body[0]
            anchor.addprevious(img)
            anchor.addprevious(brk)
        else:
            anchor = sect_ps[min(n_before - 1 + offset, len(sect_ps) - 1)].getparent()
            anchor.addnext(brk)
            anchor.addnext(img)
        offset += 1
    pdf.close()
    doc.save(str(docx_path))


def _page_image_para(doc, pdf, page_index: int, width: Emu):
    pix = pdf[page_index].get_pixmap(dpi=150)
    buf = io.BytesIO(pix.tobytes("png"))
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(buf, width=width)
    return para._p


def _page_break_para(doc):
    para = doc.add_paragraph()
    para.add_run().add_break(WD_BREAK.PAGE)
    return para._p


def doc_to_docx(doc_path: Path, out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    if sys.platform == "win32":
        try:
            return word_doc_to_docx(doc_path, out_dir)
        except ImportError:
            pass
    soffice = find_soffice()
    if not soffice:
        raise LibreOfficeNotFoundError(
            f"Cannot convert '{doc_path.name}': legacy .doc files require Microsoft Word "
            "or LibreOffice (https://www.libreoffice.org). "
            "Please save the file as .docx or .pdf and try again."
        )
    result = subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(out_dir),
            str(doc_path),
        ],
        capture_output=True,
        text=True,
        timeout=120,
    )
    converted = out_dir / (doc_path.stem + ".docx")
    if not converted.exists():
        raise RuntimeError(f"LibreOffice failed to convert {doc_path}: {result.stderr}")
    return converted


def to_docx(path: Path) -> tuple[Path, bool]:
    """Return (docx_path, is_temp_copy)."""
    ext = path.suffix.lower()
    tmp = Path(tempfile.mkdtemp(prefix="sanitizer_"))
    if ext == ".docx":
        if _has_landscape(path):
            work = tmp / path.name
            work.write_bytes(path.read_bytes())
            replace_landscape_pages_with_images(work)
            return work, True
        return path, False
    if ext == ".pdf":
        return pdf_to_docx(path, tmp), True
    if ext == ".doc":
        converted = doc_to_docx(path, tmp)
        if _has_landscape(converted):
            replace_landscape_pages_with_images(converted)
        return converted, True
    raise ValueError(f"Unsupported file type: {path}")

def _portraitize_pdf(pdf_path: Path) -> Path:
    doc = fitz.open(str(pdf_path))
    if not any(p.rect.width > p.rect.height for p in doc):
        doc.close()
        return pdf_path
    out = fitz.open()
    a4_w, a4_h = 595.27, 841.89
    for page in doc:
        r = page.rect
        if r.width > r.height:
            new = out.new_page(width=a4_w, height=a4_h)
            scale = min(a4_w / r.width, a4_h / r.height)
            w, h = r.width * scale, r.height * scale
            box = fitz.Rect((a4_w - w) / 2, (a4_h - h) / 2, (a4_w + w) / 2, (a4_h + h) / 2)
            new.show_pdf_page(box, doc, page.number)
        else:
            new = out.new_page(width=r.width, height=r.height)
            new.show_pdf_page(fitz.Rect(0, 0, r.width, r.height), doc, page.number)
    out_path = pdf_path.parent / (pdf_path.stem + ".portrait.pdf")
    out.save(str(out_path))
    out.close()
    doc.close()
    return out_path


W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _find_landscape_ranges(root):
    body = root.find(f"{W_NS}body")
    children = list(body)
    sect_ends = []
    for i, ch in enumerate(children):
        if ch.tag == f"{W_NS}sectPr" or ch.find(f"{W_NS}pPr/{W_NS}sectPr") is not None:
            sect_ends.append(i)
    landscape = []
    start = 0
    for end in sect_ends:
        sect_pr = (
            children[end].find(f"{W_NS}pPr/{W_NS}sectPr")
            if children[end].tag != f"{W_NS}sectPr"
            else children[end]
        )
        pg_sz = sect_pr.find(f"{W_NS}pgSz")
        if pg_sz is not None:
            w = int(pg_sz.get(f"{W_NS}w", "0"))
            h = int(pg_sz.get(f"{W_NS}h", "1"))
            if pg_sz.get(f"{W_NS}orient") == "landscape" or w > h:
                landscape.append((start, end, sect_pr))
        start = end + 1
    return body, children, landscape


def _has_landscape(docx_path: Path) -> bool:
    import zipfile
    from lxml import etree

    with zipfile.ZipFile(str(docx_path)) as z:
        root = etree.fromstring(z.read("word/document.xml"))
    return bool(_find_landscape_ranges(root)[2])


def replace_landscape_pages_with_images(docx_path: Path) -> int:
    from docx import Document as _Doc
    from docx.shared import Emu as _Emu

    doc2 = _Doc(str(docx_path))
    body, children, landscape = _find_landscape_ranges(doc2.element)
    if not landscape:
        return 0
    import tempfile

    with tempfile.TemporaryDirectory(prefix="sanitizer_ls_") as td:
        tdp = Path(td)
        rendered = None
        if sys.platform == "win32":
            try:
                rendered = word_docx_to_pdf(docx_path, tdp)
            except ImportError:
                rendered = None
        if rendered is None:
            soffice = find_soffice()
            if not soffice:
                return 0
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(tdp), str(docx_path)],
                capture_output=True,
                timeout=300,
            )
            if result.returncode != 0:
                return 0
            pdfs = list(tdp.glob("*.pdf"))
            if not pdfs:
                return 0
            rendered = pdfs[0]
        redacted = _redact_watermark(rendered)
        doc = fitz.open(str(redacted))
        ls_pages = [p for p in doc if p.rect.width > p.rect.height]
        if len(ls_pages) != len(landscape):
            doc.close()
            return 0
        replaced = 0
        for (start, end, sect_pr), page in zip(landscape, ls_pages):
            pg_mar = sect_pr.find(f"{W_NS}pgMar")
            left = int(pg_mar.get(f"{W_NS}left", "1440")) if pg_mar is not None else 1440
            right = int(pg_mar.get(f"{W_NS}right", "1440")) if pg_mar is not None else 1440
            usable_emu = (11906 - left - right) * 635
            pix = page.get_pixmap(dpi=150)
            import io as _io

            buf = _io.BytesIO(pix.tobytes("png"))
            para = doc2.add_paragraph()
            para.alignment = 1
            run = para.add_run()
            run.add_picture(buf, width=_Emu(usable_emu))
            new_p = para._p
            for i in range(start, end):
                body.remove(children[i])
            children[end].addprevious(new_p)
            replaced += 1
        doc.close()
        doc2.save(str(docx_path))
    return replaced
