import io
import random
import sys
from pathlib import Path

import imagehash
import re
from PIL import Image

_HASH_SIZE = 16
_MAX_DISTANCE = 10

_publisher_hashes = None
_dummy_images = None


def _assets_dir() -> Path:
    base = Path(__file__).parent / "assets" / "logos"
    if base.exists():
        return base
    meipass = Path(getattr(sys, "_MEIPASS", ".")) / "sanitizer" / "assets" / "logos"
    if meipass.exists():
        return meipass
    return base


def _load_references():
    global _publisher_hashes, _dummy_images
    if _publisher_hashes is not None:
        return
    _publisher_hashes = []
    for f in sorted((_assets_dir() / "publisher").iterdir()):
        try:
            img = Image.open(io.BytesIO(f.read_bytes())).convert("RGB")
            _publisher_hashes.append(imagehash.average_hash(img, _HASH_SIZE))
        except Exception:
            continue
    _dummy_images = []
    for f in sorted((_assets_dir() / "dummy").iterdir()):
        try:
            img = Image.open(io.BytesIO(f.read_bytes()))
            _dummy_images.append((f.suffix.lower(), img.convert("RGB")))
        except Exception:
            continue


def _is_publisher_logo(img: Image.Image) -> bool:
    _load_references()
    if not _publisher_hashes:
        return False
    h = imagehash.average_hash(img.convert("RGB"), _HASH_SIZE)
    return any(h - ref <= _MAX_DISTANCE for ref in _publisher_hashes)


def _random_dummy_bytes(target_size: tuple[int, int], fmt: str) -> bytes:
    _load_references()
    suffix, dummy = random.choice(_dummy_images)
    dw, dh = dummy.size
    tw, th = target_size
    target_ratio = tw / th if th else 1.0
    dummy_ratio = dw / dh if dh else 1.0
    if abs(target_ratio - dummy_ratio) > 0.01:
        canvas = Image.new("RGB" if fmt.lower() in ("jpeg", "jpg") else "RGBA", (tw, th),
                           (255, 255, 255) if fmt.lower() in ("jpeg", "jpg") else (0, 0, 0, 0))
        scale = min(tw / dw, th / dh)
        nw, nh = max(int(dw * scale), 1), max(int(dh * scale), 1)
        resized = dummy.resize((nw, nh))
        canvas.paste(resized, ((tw - nw) // 2, (th - nh) // 2))
        dummy = canvas
    else:
        dummy = dummy.resize((max(tw, 1), max(th, 1)))
    buf = io.BytesIO()
    if fmt.lower() in ("jpeg", "jpg"):
        dummy.save(buf, format="JPEG", quality=90)
    else:
        dummy.save(buf, format="PNG")
    return buf.getvalue()


def _image_bytes_is_logo(data: bytes) -> bool:
    try:
        img = Image.open(io.BytesIO(data))
        img.load()
        return _is_publisher_logo(img)
    except Exception:
        return False


def replace_logos_in_docx(docx_path: Path) -> int:
    import zipfile
    import re as _re2
    import xml.etree.ElementTree as _ET

    _load_references()
    if not _publisher_hashes:
        return 0
    src = zipfile.ZipFile(str(docx_path), "r")
    # Collect media targets referenced only from headers/footers so body images stay untouched.
    hf_targets: set[str] = set()
    try:
        names = src.namelist()
        hf_xmls = [n for n in names if (n.startswith("word/header") or n.startswith("word/footer")) and n.endswith(".xml")]
        for hf in hf_xmls:
            try:
                xml_bytes = src.read(hf)
                # r:embed ids in header/footer xml
                ids = _re2.findall(r'embed="([^"]+)"', xml_bytes.decode("utf8", "ignore"))
                ids += _re2.findall(r'id="([^"]+)"', xml_bytes.decode("utf8", "ignore"))
            except Exception:
                continue
            # Resolve ids via rels
            rels_candidates = [
                "word/_rels/" + hf.split("/")[-1] + ".rels",
                hf + ".rels",
                hf.replace("word/", "word/_rels/") + ".rels",
            ]
            rels_data = None
            rels_name = None
            for rn in rels_candidates:
                if rn in names:
                    rels_name = rn
                    try:
                        rels_data = src.read(rn)
                    except Exception:
                        rels_data = None
                    break
            if rels_data is None:
                continue
            try:
                root_rels = _ET.fromstring(rels_data)
                for rel in root_rels.iter():
                    rid = rel.get("Id")
                    tgt = rel.get("Target")
                    if rid in ids and tgt:
                        # Target like media/image1.png or ../media/image1.png
                        t = tgt.split("/")[-1]
                        # find full word/media/ entry
                        for n in names:
                            if n.endswith("/" + t) and n.startswith("word/media/"):
                                hf_targets.add(n)
                        # also handle direct
                        if tgt.startswith("media/"):
                            hf_targets.add("word/" + tgt)
            except Exception:
                continue
        # Fallback: if no rels found but hf xml directly references media via regex, skip
    except Exception:
        hf_targets = set()
    # If we have headers/footers but couldn't resolve, avoid touching body:
    # only replace when we have a non-empty allow-list. Empty allow-list => 0.
    if not hf_targets:
        # No header/footer image refs found; do not touch body images.
        src.close()
        return 0
    replaced = 0
    entries = []
    for info in src.infolist():
        data = src.read(info.filename)
        if info.filename in hf_targets and len(data) > 0:
            try:
                img = Image.open(io.BytesIO(data))
                img.load()
            except Exception:
                entries.append((info, data))
                continue
            if _is_publisher_logo(img):
                fmt = img.format or "PNG"
                data = _random_dummy_bytes(img.size, fmt)
                replaced += 1
        entries.append((info, data))
    src.close()
    if not replaced:
        return 0
    tmp = docx_path.with_suffix(docx_path.suffix + ".tmp")
    with zipfile.ZipFile(str(tmp), "w", zipfile.ZIP_DEFLATED) as dst:
        for info, data in entries:
            dst.writestr(info, data)
    tmp.replace(docx_path)
    return replaced


def replace_logos_in_pdf(pdf_path: Path) -> int:
    import fitz

    _load_references()
    if not _publisher_hashes:
        return 0
    doc = fitz.open(str(pdf_path))
    replaced = 0
    for page in doc:
        h = page.rect.height
        top_limit = page.rect.y0 + h * 0.15
        bottom_limit = page.rect.y1 - h * 0.10
        seen = set()
        for img in page.get_images(full=True):
            xref = img[0]
            if xref in seen:
                continue
            seen.add(xref)
            try:
                pix = fitz.Pixmap(doc, xref)
                if pix.colorspace is None:
                    continue
                if pix.n - pix.alpha > 3:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                data = pix.tobytes("png")
            except Exception:
                continue
            try:
                pil = Image.open(io.BytesIO(data))
                pil.load()
            except Exception:
                continue
            if not _is_publisher_logo(pil):
                continue
            rects = page.get_image_rects(xref)
            # Only replace if image lies in header (top 15%) or footer (bottom 10%)
            hf_rects = [r for r in rects if r.y1 <= top_limit or r.y0 >= bottom_limit]
            if not hf_rects:
                continue
            dummy = _random_dummy_bytes(pil.size, "PNG")
            for r in hf_rects:
                page.add_redact_annot(r)
                page.apply_redactions()
                page.insert_image(r, stream=dummy)
            replaced += 1
    out = None
    if replaced:
        out = pdf_path.with_suffix(".nologo.pdf")
        doc.save(str(out))
    doc.close()
    if out is not None:
        import os

        os.replace(str(out), str(pdf_path))
    return replaced

def _publisher_names() -> list[str]:
    names = []
    for f in sorted((_assets_dir() / "publisher").iterdir()):
        parts = f.stem.split("_")
        names.append(" ".join(p.upper() if len(p) <= 4 else p.capitalize() for p in parts))
    return names


def _publisher_domains() -> list[str]:
    return [name.lower().replace(" ", "") + ".com" for name in _publisher_names()]


_WORDMARK_MIN_PT = 20.0


def _replace_alias_across_runs(p_elem, alias: str, replacement: str, W: str) -> int:
    import re as _re

    t_nodes = p_elem.findall(f".//{{{W}}}t")
    if not t_nodes:
        return 0
    full = "".join(t.text or "" for t in t_nodes)
    matches = list(_re.finditer(_re.escape(alias), full, _re.IGNORECASE))
    if not matches:
        return 0
    spans = []
    pos = 0
    for t in t_nodes:
        text = t.text or ""
        spans.append((t, pos, pos + len(text)))
        pos += len(text)
    for m in reversed(matches):
        start, end = m.span()
        first = True
        for t, s, e in spans:
            if e <= start or s >= end:
                continue
            ls = max(start, s) - s
            le = min(end, e) - s
            seg = t.text or ""
            if first:
                t.text = seg[:ls] + replacement + seg[le:]
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                first = False
            else:
                t.text = seg[:ls] + seg[le:]
    return len(matches)


def replace_brand_text_in_doc(doc, W: str) -> int:
    """Replace publisher-name text. Headers/footers wordmarks get a dummy logo;
    body wordmarks become Lorum Ipsum text so body logos stay untouched."""
    import io as _io

    from docx.shared import Emu as _Emu
    from docx.text.paragraph import Paragraph as _Paragraph

    _load_references()
    if not _publisher_hashes:
        return 0
    aliases = _publisher_names()
    count = 0
    roots: list[tuple[object, bool]] = [(doc.element, False)]
    for section in doc.sections:
        for part in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            try:
                roots.append((part._element, True))
            except Exception:
                continue
    for root, is_hf in roots:
        for p_elem in list(root.iter(f"{{{W}}}p")):
            # Paragraph needs its owning element for add_run; body uses doc, HF uses part
            owner = doc if not is_hf else None
            # Resolve owner from root parent: for HF we stored part element, need its document
            # Fallback: use doc for Paragraph construction (python-docx accepts any)
            para = _Paragraph(p_elem, doc)
            full = para.text.strip()
            hit_alias = None
            for alias in aliases:
                if alias.lower() in full.lower():
                    hit_alias = alias
                    break
            for domain in _publisher_domains():
                if domain.lower() in full.lower():
                    count += _replace_alias_across_runs(p_elem, domain, "lorumipsum.com", W)
            if hit_alias is None:
                continue
            if len(full) <= len(hit_alias) + 10:
                if is_hf:
                    for run in list(para.runs):
                        run.text = ""
                    size_pt = 0
                    for r_el in p_elem.findall(f".//{{{W}}}sz"):
                        try:
                            size_pt = max(size_pt, int(r_el.get(f"{{{W}}}val", "0")) / 2)
                        except (TypeError, ValueError):
                            continue
                    if size_pt <= 0:
                        size_pt = 14.0
                    width_pt = min(size_pt * 0.6 * max(len(hit_alias), 1), 550)
                    height_pt = min(size_pt * 1.3, 170)
                    try:
                        run = para.add_run()
                        buf = _io.BytesIO(
                            _random_dummy_bytes((int(width_pt * 4), int(height_pt * 4)), "PNG")
                        )
                        run.add_picture(
                            buf,
                            width=_Emu(int(width_pt * 12700)),
                            height=_Emu(int(height_pt * 12700)),
                        )
                    except Exception:
                        pass
                    count += 1
                else:
                    count += _replace_alias_across_runs(p_elem, hit_alias, "Lorum Ipsum", W)
            else:
                count += _replace_alias_across_runs(p_elem, hit_alias, "Lorum Ipsum", W)
    return count


def replace_brand_text_in_pdf(pdf_path: Path) -> int:
    import fitz

    _load_references()
    if not _publisher_hashes:
        return 0
    aliases = _publisher_names()
    doc = fitz.open(str(pdf_path))
    count = 0
    for page in doc:
        edits = []
        d = page.get_text("dict")
        targets = [(a, "Lorum Ipsum") for a in aliases] + [
            (d, "lorumipsum.com") for d in _publisher_domains()
        ]
        for block in d["blocks"]:
            for line in block.get("lines", []):
                for span in line["spans"]:
                    text = span["text"]
                    for alias, repl in targets:
                        if alias.lower() not in text.lower():
                            continue
                        subs = page.search_for(alias, clip=fitz.Rect(span["bbox"]))
                        for r in subs:
                            edits.append((r, span["size"], span["color"], span["bbox"], repl))
        if not edits:
            continue
        edited_rects = [e[0] for e in edits]
        for alias, repl in targets:
            for r in page.search_for(alias):
                if any(r.intersects(er) for er in edited_rects):
                    continue
                edits.append((r, 8.0, 0, r, repl))
                edited_rects.append(r)
        for r, size, color, _, _ in edits:
            page.add_redact_annot(r)
        page.apply_redactions()
        for r, size, color, _, repl in edits:
            if size >= _WORDMARK_MIN_PT and repl == "Lorum Ipsum":
                dummy = _random_dummy_bytes((int(r.width * 2), int(r.height * 2)), "PNG")
                fit_r = fitz.Rect(r.x0, r.y0, min(r.x0 + r.height * 2.5, page.rect.width), r.y1)
                page.insert_image(fit_r, stream=dummy)
            else:
                c = color
                rgb = ((c >> 16 & 255) / 255, (c >> 8 & 255) / 255, (c & 255) / 255)
                page.insert_text(fitz.Point(r.x0, r.y1 - 1), repl,
                                 fontsize=size, fontname="helv", color=rgb)
            count += 1
    if count:
        out = pdf_path.with_suffix(".nobrand.pdf")
        doc.save(str(out))
        doc.close()
        import os

        os.replace(str(out), str(pdf_path))
    else:
        doc.close()
    return count

_brand_hues_cache = None


def _brand_hsv_colors():
    global _brand_hues_cache
    if _brand_hues_cache is not None:
        return _brand_hues_cache
    import colorsys

    colors = []
    for f in sorted((_assets_dir() / "publisher").iterdir()):
        try:
            img = Image.open(io.BytesIO(f.read_bytes())).convert("RGBA").resize((64, 64))
        except Exception:
            continue
        for px in list(img.getdata()):
            r, g, b, a = px[0], px[1], px[2], px[3]
            if a < 128:
                continue
            h, s, v = colorsys.rgb_to_hsv(r / 255, g / 255, b / 255)
            if s >= 0.3 and v >= 0.15:
                colors.append(h * 360)
    _brand_hues_cache = colors
    return colors


def _hue_matches(hex_color: str, brand_hues) -> bool:
    import colorsys

    m = re.fullmatch(r"#?([0-9a-fA-F]{6})", hex_color.strip())
    if not m:
        return False
    r = int(m.group(1)[0:2], 16) / 255
    g = int(m.group(1)[2:4], 16) / 255
    b = int(m.group(1)[4:6], 16) / 255
    h, s, v = colorsys.rgb_to_hsv(r, g, b)
    if s < 0.3 or v < 0.15:
        return False
    hue = h * 360
    return any(min(abs(hue - bh), 360 - abs(hue - bh)) <= 22 for bh in brand_hues)


def replace_vector_logos_in_doc(doc, W: str) -> int:
    """Replace tracked-brand vector logos with random dummies. Headers/footers
    are always scanned; body only at document corners (first/last page cover)
    so page-5 charts stay untouched."""
    import io as _io
    import re as _re
    import bisect as _bisect_mod

    from docx.shared import Emu as _Emu
    from docx.text.paragraph import Paragraph as _Paragraph

    _load_references()
    if not _publisher_hashes:
        return 0
    brand_hues = _brand_hsv_colors()
    root = doc.element

    story_roots: list[tuple[object, object, bool]] = []
    for section in doc.sections:
        for part in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            try:
                story_roots.append((part._element, part, True))
            except Exception:
                continue
    # Body only for first/last page corners (covers) — not mid-document like page 5
    story_roots.append((root, doc, False))
    if not story_roots:
        return 0

    def _collect(root, is_hf):
        cands = []
        for pict in list(root.iter(f"{{{W}}}pict")):
            if pict.find(f".//{{{V}}}imagedata") is not None:
                continue
            if pict.find(f".//{{{W}}}txbxContent") is not None:
                continue
            colors = [
                el.get("fillcolor", "") for el in pict.iter() if el.get("fillcolor")
            ] + [el.get("strokecolor", "") for el in pict.iter() if el.get("strokecolor")]
            color_hit = any(_hue_matches(cl, brand_hues) for cl in colors)
            if not color_hit:
                if not is_hf:
                    continue
                fills = [cl.lower() for cl in colors if cl]
                if not fills or not all(f in ("#000000", "black", "#000") for f in fills):
                    continue
            style = ""
            for shape in pict.iter():
                if shape.get("style"):
                    style = shape.get("style")
                    break
            m_w = _re.search(r"width:([\d.]+)pt", style)
            m_h = _re.search(r"height:([\d.]+)pt", style)
            m_l = _re.search(r"margin-left:([-\d.]+)pt", style)
            m_t = _re.search(r"margin-top:([-\d.]+)pt", style)
            if not (m_w and m_h):
                continue
            w_pt = float(m_w.group(1))
            h_pt = float(m_h.group(1))
            if not (4 <= h_pt <= 120 and 5 <= w_pt <= 550):
                continue
            if not is_hf and not (8 <= h_pt <= 170 and 15 <= w_pt <= 550):
                continue
            shape_count = sum(
                1 for el in pict.iter() if el.tag in (f"{{{V}}}shape", f"{{{V}}}group")
            )
            if shape_count > (400 if is_hf else 120):
                continue
            run_el = pict.getparent()
            while run_el is not None and run_el.tag != f"{{{W}}}r":
                run_el = run_el.getparent()
            p_el = run_el.getparent() if run_el is not None else None
            if p_el is None:
                continue
            cands.append({
                "pict": pict, "run": run_el, "p": p_el,
                "ml": float(m_l.group(1)) if m_l else 0.0,
                "mt": float(m_t.group(1)) if m_t else 0.0,
                "w": w_pt, "h": h_pt,
            })
        return cands

    # Document corners = very first child (cover) and very last section start
    body = root.find(f"{{{W}}}body")
    body_children = list(body) if body is not None else []
    child_index = {id(ch): i for i, ch in enumerate(body_children)}
    boundaries = [
        i for i, ch in enumerate(body_children)
        if ch.tag == f"{{{W}}}sectPr" or ch.find(f"{{{W}}}pPr/{{{W}}}sectPr") is not None
    ]
    section_starts = [0] + [b + 1 for b in boundaries[:-1]] if boundaries else [0]
    last_start = section_starts[-1] if section_starts else 0

    def _is_document_corner(p_el) -> bool:
        top = p_el
        while top.getparent() is not None and top.getparent() is not body:
            top = top.getparent()
        idx = child_index.get(id(top))
        if idx is None:
            return False
        # First page cover: first 2 children only (not all before first_text which includes page 5)
        if idx <= 2:
            return True
        # Last page cover: within first 3 children of last section
        return idx >= last_start and idx <= last_start + 3

    replaced = 0
    for root, parent, is_hf in story_roots:
        cands = _collect(root, is_hf)
        if not cands:
            continue
        if not is_hf:
            cands = [cd for cd in cands if _is_document_corner(cd["p"])]
            if not cands:
                continue
        cands = sorted(cands, key=lambda cd: (id(cd["p"]), cd["ml"]))
        used = set()
        for i, cd in enumerate(cands):
            if i in used:
                continue
            cluster = [cd]
            used.add(i)
            for j in range(i + 1, len(cands)):
                if j in used or cands[j]["p"] is not cd["p"]:
                    continue
                last = cluster[-1]
                gap = cands[j]["ml"] - (last["ml"] + last["w"])
                voverlap = abs(cands[j]["mt"] - last["mt"]) < max(cands[j]["h"], last["h"]) + 6
                if gap > 30 or not voverlap:
                    continue
                cluster.append(cands[j])
                used.add(j)
            cluster.sort(key=lambda cd: cd["ml"])
            x0 = min(cd["ml"] for cd in cluster)
            x1 = max(cd["ml"] + cd["w"] for cd in cluster)
            y0 = min(cd["mt"] for cd in cluster)
            y1 = max(cd["mt"] + cd["h"] for cd in cluster)
            w_pt = max(x1 - x0, 15)
            h_pt = max(y1 - y0, 8)
            anchor_p = cluster[0]["p"]
            para = _Paragraph(anchor_p, parent)
            new_run = para.add_run()
            buf = _io.BytesIO(_random_dummy_bytes((int(w_pt * 4), int(h_pt * 4)), "PNG"))
            new_run.add_picture(buf, width=_Emu(int(w_pt * 12700)), height=_Emu(int(h_pt * 12700)))
            anchor_p.remove(new_run._r)
            cluster[0]["run"].addnext(new_run._r)
            spacing = anchor_p.find(f"{{{W}}}pPr/{{{W}}}spacing")
            if spacing is not None:
                spacing.set(f"{{{W}}}line", str(int(h_pt * 20 * 1.15)))
                spacing.set(f"{{{W}}}lineRule", "atLeast")
            for cd in cluster:
                if cd["run"].getparent() is anchor_p:
                    anchor_p.remove(cd["run"])
            replaced += 1
    return replaced





V = "urn:schemas-microsoft-com:vml"
